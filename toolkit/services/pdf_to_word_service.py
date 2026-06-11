from io import BytesIO

import fitz

from ..utils.file_helpers import get_temp_path, cleanup_files

try:
    from pdf2docx import Converter
except ImportError:
    Converter = None


def has_text_layer(pdf_path, sample_pages=10):
    doc = fitz.open(pdf_path)
    try:
        for index, page in enumerate(doc):
            if index >= sample_pages:
                break
            if page.get_text().strip():
                return True
        return False
    finally:
        doc.close()


def pdf_to_word(pdf_path):
    """Convert a text-based PDF to .docx. Returns BytesIO."""
    if not has_text_layer(pdf_path):
        raise ValueError(
            "This looks like a scanned PDF with no text layer. "
            "Run it through the OCR PDF tool first, then convert."
        )

    if Converter is not None:
        out_path = get_temp_path(".docx")
        cv = Converter(pdf_path)
        try:
            cv.convert(out_path)
        finally:
            cv.close()
        with open(out_path, "rb") as fh:
            buf = BytesIO(fh.read())
        cleanup_files(out_path)
        return buf

    return _pdf_to_word_basic(pdf_path)


def _pdf_to_word_basic(pdf_path):
    """Fallback when pdf2docx is unavailable: text blocks → paragraphs."""
    from docx import Document
    from docx.shared import Pt

    doc = fitz.open(pdf_path)
    try:
        word_doc = Document()
        for page_index, page in enumerate(doc):
            if page_index > 0:
                word_doc.add_page_break()
            for block in page.get_text("dict").get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    paragraph = word_doc.add_paragraph()
                    for span in spans:
                        text = span.get("text") or ""
                        if not text:
                            continue
                        run = paragraph.add_run(text)
                        run.font.size = Pt(round(span.get("size") or 11))
                        run.bold = bool(span.get("flags", 0) & 16)
                        run.italic = bool(span.get("flags", 0) & 2)

        out = BytesIO()
        word_doc.save(out)
        out.seek(0)
        return out
    finally:
        doc.close()
